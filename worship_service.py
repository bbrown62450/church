#!/usr/bin/env python3
"""
Worship service generator: hymn suggestions by scripture, OpenAI liturgy,
and Word document export.
"""

import logging
import os
import re
from typing import Dict, Any, List, Optional
from io import BytesIO
from urllib.parse import urlparse, urlunparse

from dotenv import load_dotenv
from notion_hymns import NotionHymnsDB
from hymn_utils import get_property_value

load_dotenv()

logger = logging.getLogger(__name__)

# Cache for resolved Hymnary audio URLs (number -> url) to avoid re-fetching
_hymnary_audio_resolve_cache: Dict[int, Optional[str]] = {}

# Optional imports for docx and openai
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


def _add_leader_people_paragraph(doc, text: str) -> None:
    """Add paragraph(s): each Leader: (normal), each People: (bold). Supports multiple Leader/People pairs."""
    if not text or not text.strip():
        return
    # Find all "Leader:" and "People:" in order (case-insensitive)
    pattern = re.compile(r"\b(Leader|People):\s*", re.IGNORECASE)
    pos = 0
    parts = []
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append(("", text[pos : m.start()].strip()))  # preamble if any
        role = "Leader" if m.group(1).lower() == "leader" else "People"
        end = pattern.search(text, m.end())
        content_end = end.start() if end else len(text)
        content = text[m.end() : content_end].strip()
        parts.append((role, content))
        pos = content_end
    if not parts:
        doc.add_paragraph(text)
        return
    for role, content in parts:
        if not content and role == "":
            continue
        if role == "People":
            p = doc.add_paragraph()
            p.add_run("People: ")
            r = p.add_run(content)
            r.bold = True
        else:
            line = ("Leader: " + content) if role == "Leader" else content
            if line:
                doc.add_paragraph(line)


def _add_communion_liturgy(doc) -> None:
    """Add The Sacrament of the Lord's Supper liturgy (invitation, great thanksgiving, etc.)."""
    doc.add_paragraph("The Sacrament of the Lord's Supper", style="Heading 1")
    doc.add_paragraph()

    doc.add_paragraph("Invitation to the Table", style="Heading 2")
    doc.add_paragraph(
        "This is the table of our Lord Jesus Christ. It is not a reward for the righteous, "
        "but nourishment for those who hunger; not a prize for the strong, but grace for those who are weary. "
        "Here, blessing is not earned but received. All who seek to walk humbly with God, and who trust in God's mercy, "
        "are welcome at this table."
    )
    doc.add_paragraph()

    doc.add_paragraph("Great Thanksgiving", style="Heading 2")
    doc.add_paragraph("The Lord be with you.")
    p = doc.add_paragraph()
    p.add_run("And also with you.").bold = True
    doc.add_paragraph("Lift up your hearts.")
    p = doc.add_paragraph()
    p.add_run("We lift them up to the Lord.").bold = True
    doc.add_paragraph("Let us give thanks to the Lord our God.")
    p = doc.add_paragraph()
    p.add_run("It is right to give our thanks and praise.").bold = True
    doc.add_paragraph(
        "It is truly right and our greatest joy to give you thanks and praise, O God, "
        "creator of heaven and earth, for you have made us and all things, and in your love you hold us in life. "
        "And so we join the everlasting song:"
    )
    p = doc.add_paragraph()
    p.add_run(
        "Holy, holy, holy Lord, God of power and might, heaven and earth are full of your glory. "
        "Hosanna in the highest. Blessed is the one who comes in the name of the Lord. Hosanna in the highest."
    ).bold = True
    doc.add_paragraph(
        "You are holy, O God of majesty, and blessed is Jesus Christ, your Son, our Lord, "
        "who by his life, death, and resurrection has reconciled the world to you. On the night in which he gave himself up "
        "he took bread, gave thanks, broke it, and gave it to his disciples. And likewise the cup after supper. "
        "Remembering his death and resurrection, we offer ourselves in praise and thanksgiving. Therefore we proclaim the mystery of faith:"
    )
    doc.add_paragraph("Christ has died.")
    doc.add_paragraph("Christ is risen.")
    doc.add_paragraph("Christ will come again.")
    doc.add_paragraph()

    doc.add_paragraph("Words of Institution", style="Heading 2")
    doc.add_paragraph("[Words of institution as printed or as used.]")
    doc.add_paragraph()

    doc.add_paragraph("The Lord's Prayer", style="Heading 2")
    doc.add_paragraph("[The Lord's Prayer as printed.]")
    doc.add_paragraph()

    doc.add_paragraph("Breaking of the Bread and Communion", style="Heading 2")
    doc.add_paragraph(
        "The bread that we break is a sharing in the body of Christ. "
        "The cup that we bless is a sharing in the blood of Christ. Come, for all is ready."
    )
    doc.add_paragraph()

    doc.add_paragraph("Prayer After Communion", style="Heading 2")
    doc.add_paragraph(
        "Gracious God, we give you thanks that you have fed us at this table of grace, "
        "strengthening us not to win our lives, but to live them faithfully. Send us out to do justice, "
        "to love kindness, and to walk humbly with you, bearing your blessing into a world still hungry for hope, "
        "through Jesus Christ our Lord. Amen."
    )
    doc.add_paragraph()


def _add_assurance_paragraph(doc, leader_text: str) -> None:
    """Add Assurance: Leader line then 'People: Thanks be to God! Amen.' in bold."""
    leader_clean = (leader_text or "").strip()
    if leader_clean.startswith("Leader:"):
        leader_clean = leader_clean[7:].strip()
    if leader_clean:
        doc.add_paragraph("Leader: " + leader_clean)
    # Always add the congregational response
    p = doc.add_paragraph()
    r = p.add_run("People: Thanks be to God! Amen.")
    r.bold = True


# Common Bible book abbreviations (full name -> variants to try)
_BOOK_ABBREVS: Dict[str, List[str]] = {
    "genesis": ["gen", "ge", "gn"],
    "exodus": ["exod", "ex"],
    "leviticus": ["lev"],
    "numbers": ["num", "nm"],
    "deuteronomy": ["deut", "dt"],
    "joshua": ["josh", "jos"],
    "judges": ["judg", "jdg"],
    "ruth": [],
    "1 samuel": ["1 sam", "1 samuel"],
    "2 samuel": ["2 sam", "2 samuel"],
    "1 kings": ["1 kgs", "1 kings"],
    "2 kings": ["2 kgs", "2 kings"],
    "1 chronicles": ["1 chr", "1 chron"],
    "2 chronicles": ["2 chr", "2 chron"],
    "ezra": [],
    "nehemiah": ["neh"],
    "esther": ["esth", "est"],
    "job": [],
    "psalm": ["ps", "psa", "psalms"],
    "psalms": ["ps", "psa", "psalm"],
    "proverbs": ["prov", "prv"],
    "ecclesiastes": ["eccl", "ecc"],
    "song of solomon": ["song", "sos", "canticles"],
    "song of songs": ["song", "sos"],
    "isaiah": ["isa", "is"],
    "jeremiah": ["jer", "jr"],
    "lamentations": ["lam"],
    "ezekiel": ["ezek", "ezk"],
    "daniel": ["dan", "dnl"],
    "hosea": ["hos"],
    "joel": [],
    "amos": [],
    "obadiah": ["obad", "ob"],
    "jonah": ["jon"],
    "micah": ["mic"],
    "nahum": ["nah"],
    "habakkuk": ["hab"],
    "zephaniah": ["zeph", "zep"],
    "haggai": ["hag"],
    "zechariah": ["zech", "zec"],
    "malachi": ["mal"],
    "matthew": ["matt", "mt"],
    "mark": ["mk", "mrk"],
    "luke": ["lk", "luk"],
    "john": ["jn", "jhn", "joh"],
    "acts": [],
    "romans": ["rom", "rm"],
    "1 corinthians": ["1 cor", "1 corinthians"],
    "2 corinthians": ["2 cor", "2 corinthians"],
    "galatians": ["gal"],
    "ephesians": ["eph"],
    "philippians": ["phil", "php"],
    "colossians": ["col"],
    "1 thessalonians": ["1 thess", "1 thes"],
    "2 thessalonians": ["2 thess", "2 thes"],
    "1 timothy": ["1 tim", "1 timothy"],
    "2 timothy": ["2 tim", "2 timothy"],
    "titus": ["tit"],
    "philemon": ["phlm", "phm"],
    "hebrews": ["heb"],
    "james": ["jas", "jm"],
    "1 peter": ["1 pet", "1 peter"],
    "2 peter": ["2 pet", "2 peter"],
    "1 john": ["1 jn", "1 jhn"],
    "2 john": ["2 jn", "2 jhn"],
    "3 john": ["3 jn", "3 jhn"],
    "jude": ["jud"],
    "revelation": ["rev", "rv"],
}


def _scripture_search_variants(ref: str) -> List[str]:
    """
    Generate fuzzy search variants from a scripture reference.
    E.g. "Genesis 12:1-4a" -> ["Genesis 12:1-4a", "Genesis 12:1-4", "Genesis 12:1", "Genesis 12", "Gen 12", ...]
    """
    ref = ref.strip()
    if not ref:
        return []
    variants = [ref]
    ref_lower = ref.lower()

    # Extract book + chapter + verse
    # Match: "Genesis 12:1-4a" or "Psalm 121" or "Romans 4:1-5, 13-17"
    m = re.match(r"^(.+?)\s+(\d+)(?::(.+))?$", ref, re.IGNORECASE)
    if m:
        book_part = m.group(1).strip()
        chapter = m.group(2)
        verse_part = m.group(3) or ""

        # Book + chapter (e.g. "Genesis 12")
        bc = f"{book_part} {chapter}"
        variants.append(bc)

        if verse_part:
            # Strip "a" or "b" suffix (e.g. "1-4a" -> "1-4")
            verse_clean = re.sub(r"([\d\-]+)[ab]\b", r"\1", verse_part, flags=re.I).strip()
            if verse_clean:
                bcv = f"{book_part} {chapter}:{verse_clean}"
                variants.append(bcv)
            # First verse only (e.g. "Genesis 12:1")
            first_verse = re.match(r"(\d+)", verse_part)
            if first_verse:
                variants.append(f"{book_part} {chapter}:{first_verse.group(1)}")

    # Abbreviation variants (e.g. "Gen 12", "Matt 17")
    for full_book, abbrevs in _BOOK_ABBREVS.items():
        if ref_lower.startswith(full_book + " "):
            rest = ref[len(full_book) + 1:].strip()
            for ab in abbrevs[:2]:
                if rest:
                    variants.append(f"{ab} {rest}")

    # Dedupe while preserving order
    seen: set = set()
    out = []
    for v in variants:
        vc = v.strip()
        if vc and vc.lower() not in seen:
            seen.add(vc.lower())
            out.append(vc)
    return out


def hymns_by_scripture(
    db: NotionHymnsDB,
    scripture_ref: str,
    limit: int = 50,
    all_hymns: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    """
    Return hymns whose Scripture References contain the given reference.
    Uses fuzzy matching: tries multiple variants (book+chapter, abbreviations, etc.)
    to find more hymns. Handles " or " for alternative gospel choices.

    If *all_hymns* is provided the search runs entirely in-memory (zero API calls).
    Otherwise falls back to Notion search API + list_hymns.
    """
    ref_clean = scripture_ref.strip()
    if not ref_clean:
        return []

    refs_to_try = [ref_clean]
    if " or " in ref_clean:
        refs_to_try = [p.strip() for p in ref_clean.split(" or ") if p.strip()]

    variant_set: set = set()
    for ref in refs_to_try:
        variant_set.update(v.lower() for v in _scripture_search_variants(ref))
    variant_set.add(ref_clean.lower())
    for r in refs_to_try:
        variant_set.add(r.lower())

    seen_ids: set = set()
    results: List[Dict[str, Any]] = []

    hymn_pool = all_hymns if all_hymns is not None else None

    if hymn_pool is None:
        # No cached data — use Notion search API (original path)
        for ref in refs_to_try:
            variants = _scripture_search_variants(ref)
            if not variants:
                variants = [ref]
            variants = list(dict.fromkeys(variants))
            for variant in variants:
                if len(results) >= limit:
                    break
                try:
                    batch = db.search_hymns(
                        filter_property="Scripture References",
                        filter_value=variant,
                    )
                    for h in batch:
                        if h["id"] not in seen_ids:
                            seen_ids.add(h["id"])
                            results.append(h)
                except Exception:
                    pass
        if len(results) < 15:
            hymn_pool = db.list_hymns()

    if hymn_pool is not None:
        for h in hymn_pool:
            if h["id"] in seen_ids:
                continue
            scripture = get_property_value(h, "Scripture References")
            if not scripture:
                continue
            scripture_lower = scripture.lower()
            for v in variant_set:
                if v and v in scripture_lower:
                    results.append(h)
                    seen_ids.add(h["id"])
                    break
            if len(results) >= limit:
                break

    return results[:limit]


# Theme keywords for role-based hymn filtering
_OPENING_THEMES = {"gathering", "opening", "call to worship", "invitation", "welcome", "entrance"}
_CLOSING_THEMES = {"joy", "rejoice", "sending", "benediction", "mission", "dismissal", "praise", "thanksgiving"}


def _hymn_matches_theme(hymn: Dict[str, Any], theme_set: set) -> bool:
    """True if hymn's Theme property contains any of the theme keywords."""
    themes = get_property_value(hymn, "Theme")
    if not themes:
        return False
    if isinstance(themes, str):
        themes = [themes]
    themes_lower = " ".join(t.lower() for t in themes)
    return any(kw in themes_lower for kw in theme_set)


def suggest_hymns_for_service(
    *,
    db: NotionHymnsDB,
    occasion: str,
    scriptures: List[str],
    selected_nt_ref: Optional[str] = None,
    scripture_full_texts: Optional[Dict[str, str]] = None,
    scripture_text_fetcher: Optional[Any] = None,
    api_key: Optional[str] = None,
    limit_per_slot: int = 5,
    progress_callback: Optional[Any] = None,
    all_hymns: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, List[Dict[str, Any]]]:
    """
    Use AI to suggest hymns for opening, response (after sermon), and closing.
    - Opening: gathering/opening hymns
    - Response: hymns that match scripture themes (especially NT reading)
    - Closing: joyful, upbeat hymns

    If *all_hymns* is provided, uses that cached list instead of calling db.list_hymns()
    (avoids redundant Notion API calls).

    Returns {"opening": [...], "response": [...], "closing": [...]} with hymn info dicts.
    """
    scripture_full_texts = scripture_full_texts or {}
    client = None
    if OpenAI:
        key = api_key or os.getenv("OPENAI_API_KEY")
        if key:
            client = OpenAI(api_key=key)

    if not client:
        return {"opening": [], "response": [], "closing": []}

    def _progress(msg: str, pct: float) -> None:
        if progress_callback:
            progress_callback(msg, pct)

    # Get NT scripture text for theme extraction
    _progress("Preparing scripture text…", 0.02)
    nt_text = ""
    if selected_nt_ref:
        nt_text = scripture_full_texts.get(selected_nt_ref) or ""
        if not nt_text and scripture_text_fetcher:
            try:
                nt_text = scripture_text_fetcher(selected_nt_ref) or ""
            except Exception:
                pass
        if nt_text and "[Could not load text]" in nt_text:
            nt_text = ""
    if not nt_text and scriptures:
        # Fallback: use first scripture ref that looks NT
        for ref in scriptures:
            if " or " in ref:
                for part in ref.split(" or "):
                    if any(b in part.lower() for b in ["matthew", "mark", "luke", "john", "acts", "romans", "corinthians", "galatians", "ephesians", "philippians", "colossians", "thessalonians", "timothy", "titus", "philemon", "hebrews", "peter", "jude", "revelation"]):
                        nt_text = scripture_full_texts.get(part.strip()) or ""
                        if nt_text:
                            break
            if nt_text:
                break

    # Gather candidate hymns
    if all_hymns is None:
        _progress("Loading hymn list from Notion…", 0.1)
        all_hymns = db.list_hymns()
    else:
        _progress("Using cached hymn list…", 0.1)
    scripture_hymns = []
    total_refs = sum(1 + ref.count(" or ") for ref in scriptures) if scriptures else 0
    ref_idx = 0
    for ref in scriptures:
        if " or " in ref:
            for part in ref.split(" or "):
                _progress(f"Searching hymns for {part.strip()}…", 0.15 + 0.25 * (ref_idx / max(total_refs, 1)))
                scripture_hymns.extend(hymns_by_scripture(db, part.strip(), limit=30, all_hymns=all_hymns))
                ref_idx += 1
        else:
            _progress(f"Searching hymns for {ref}…", 0.15 + 0.25 * (ref_idx / max(total_refs, 1)))
            scripture_hymns.extend(hymns_by_scripture(db, ref, limit=30, all_hymns=all_hymns))
            ref_idx += 1
    seen = set()
    scripture_hymns = [h for h in scripture_hymns if h["id"] not in seen and not seen.add(h["id"])]

    opening_candidates = [h for h in all_hymns if _hymn_matches_theme(h, _OPENING_THEMES)]
    if not opening_candidates:
        opening_candidates = all_hymns[:80]

    closing_candidates = [h for h in all_hymns if _hymn_matches_theme(h, _CLOSING_THEMES)]
    if not closing_candidates:
        closing_candidates = all_hymns[:80]

    response_candidates = scripture_hymns if scripture_hymns else all_hymns[:80]

    _progress("Building prompt for AI…", 0.45)
    def _hymn_summary(h: Dict) -> str:
        title = get_property_value(h, "Hymn Title") or "Unknown"
        num = get_property_value(h, "Hymn Number")
        themes = get_property_value(h, "Theme")
        themes_str = ", ".join(themes) if isinstance(themes, list) else (themes or "")
        script = get_property_value(h, "Scripture References") or ""
        return f"- {title} (#{num})" + (f" [themes: {themes_str}]" if themes_str else "") + (f" [scripture: {script[:60]}...]" if len(script) > 60 else f" [scripture: {script}]" if script else "")

    opening_list = "\n".join(_hymn_summary(h) for h in opening_candidates[:60])
    response_list = "\n".join(_hymn_summary(h) for h in response_candidates[:60])
    closing_list = "\n".join(_hymn_summary(h) for h in closing_candidates[:60])

    scripture_refs_str = "\n".join(f"- {s}" for s in scriptures) if scriptures else "None"
    nt_preview = (nt_text[:1500] + "...") if len(nt_text) > 1500 else nt_text if nt_text else "(no text loaded)"

    prompt = f"""You are helping plan a worship service. Select hymns for three slots.

OCCASION: {occasion}
SCRIPTURE READINGS: {scripture_refs_str}
NEW TESTAMENT READING (for response hymn): {selected_nt_ref or "Not specified"}
NT PASSAGE TEXT (excerpt): {nt_preview}

ROLE REQUIREMENTS:
- OPENING: Must be a gathering/opening hymn—something that invites people into worship, calls them to praise, or welcomes them. NOT a hymn focused on the sermon theme.
- RESPONSE (after sermon): Must connect to the scripture, especially the New Testament reading. Match themes in the passage (e.g. transfiguration, Lent, grace, faith, etc.).
- CLOSING: Must be joyful, upbeat, or sending—something that sends people out with hope and praise. NOT somber or reflective.

CANDIDATE HYMNS:

OPENING CANDIDATES (prefer gathering/opening hymns):
{opening_list}

RESPONSE CANDIDATES (prefer scripture-linked hymns):
{response_list}

CLOSING CANDIDATES (prefer joyful/sending hymns):
{closing_list}

Respond with a JSON object only, no other text:
{{"opening": ["Exact Hymn Title 1", "Exact Hymn Title 2", ...], "response": ["Exact Hymn Title 1", ...], "closing": ["Exact Hymn Title 1", ...]}}

Pick {limit_per_slot} hymns per slot. Use the EXACT titles from the lists above."""

    _progress("Calling AI to select hymns…", 0.55)
    try:
        resp = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL", "gpt-3.5-turbo"),
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        content = (resp.choices[0].message.content or "").strip()
        # Extract JSON (handle markdown code blocks)
        if "```" in content:
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        content = content.strip()
        import json
        data = json.loads(content)
    except Exception as e:
        logger.warning("suggest_hymns_for_service failed: %s", e)
        return {"opening": [], "response": [], "closing": []}

    _progress("Applying suggestions…", 0.9)
    logger.info("AI raw response: %s", data)
    title_to_hymn = {}
    for h in all_hymns:
        t = get_property_value(h, "Hymn Title")
        if t:
            title_to_hymn[t.strip().lower()] = h

    def _resolve(slot: str) -> List[Dict[str, Any]]:
        titles = data.get(slot, [])
        out = []
        for t in titles:
            if isinstance(t, str) and t.strip():
                key = t.strip().lower()
                if key in title_to_hymn:
                    info = hymn_display_info(title_to_hymn[key])
                    out.append(info)
                    logger.info("Resolved %s: %r -> exact match", slot, t)
                else:
                    matched = False
                    for k, h in title_to_hymn.items():
                        if key in k or k in key:
                            info = hymn_display_info(h)
                            out.append(info)
                            logger.info("Resolved %s: %r -> fuzzy match %r", slot, t, k)
                            matched = True
                            break
                    if not matched:
                        logger.warning("Resolve %s: %r -> NO MATCH in %d hymns", slot, t, len(title_to_hymn))
        logger.info("Resolved %s: %d of %d titles matched", slot, len(out), len(titles))
        return out

    return {
        "opening": _resolve("opening"),
        "response": _resolve("response"),
        "closing": _resolve("closing"),
    }


def hymn_display_info(hymn: Dict[str, Any], *, resolve_audio: bool = False) -> Dict[str, Any]:
    """Extract title, number, and link for display/export.
    If resolve_audio is True, fetches the hymn page to get the real MP3 URL (used for scripture list players).
    """
    title = get_property_value(hymn, "Hymn Title") or "Unknown"
    number = get_property_value(hymn, "Hymn Number")
    link = get_property_value(hymn, "Hymnary.org Link")
    if number is not None:
        audio_url = (
            resolve_hymnary_audio_url(number, title)
            if resolve_audio
            else _hymnary_audio_url(number, title)
        )
    else:
        audio_url = None
    logger.debug(
        "hymn_display_info hymn_id=%s number=%s title=%r audio_url=%s",
        hymn.get("id"),
        number,
        title,
        audio_url,
    )
    return {
        "title": title,
        "number": number,
        "link": link,
        "audio_url": audio_url,
    }


def _hymnary_audio_url(number: Optional[int], title: str) -> Optional[str]:
    """
    Build a possible Hymnary.org GG2013 audio MP3 URL. Pattern from hymnary:
    .../hymnary/audio/GG2013/{number:03d}-{slug}.mp3
    Slug is a short lowercase version of the title (spaces as %20). Not all hymns have audio.
    Some hymns use a different CDN (e.g. 150282) and slug format; use resolve_hymnary_audio_url() for those.
    """
    if number is None:
        logger.debug("_hymnary_audio_url number=None title=%r -> None", title)
        return None
    slug = re.sub(r"[^\w\s]", "", (title or "").lower())
    slug = re.sub(r"\s+", "%20", slug.strip())[:30]
    if not slug:
        slug = str(number)
    num_str = f"{int(number):03d}"
    url = f"https://hymnary.org/media/fetch/148542/hymnary/audio/GG2013/{num_str}-{slug}.mp3"
    logger.debug("_hymnary_audio_url number=%s title=%r -> %s", number, title, url)
    return url


def resolve_hymnary_audio_url(number: Optional[int], title: str) -> Optional[str]:
    """
    Resolve the real GG2013 audio URL by fetching the hymn page and parsing the MP3 link.
    Hymnary uses different CDN IDs and slug formats (e.g. 191 uses 150282 and WeHaveComeAt_accomp),
    so the constructed URL can point at the wrong file. Results are cached by hymn number.
    Never raises: on any error returns the constructed fallback URL so the hymn list still renders.
    """
    if number is None:
        return None
    num = int(number)
    if num in _hymnary_audio_resolve_cache:
        return _hymnary_audio_resolve_cache[num]
    fallback = _hymnary_audio_url(number, title)
    try:
        import httpx
        from bs4 import BeautifulSoup
    except ImportError:
        _hymnary_audio_resolve_cache[num] = fallback
        return fallback
    page_url = f"https://hymnary.org/hymn/GG2013/{num}"
    try:
        with httpx.Client(follow_redirects=True, timeout=8.0) as client:
            r = client.get(page_url)
            r.raise_for_status()
    except Exception as e:
        logger.debug("resolve_hymnary_audio_url fetch %s: %s", page_url, e)
        _hymnary_audio_resolve_cache[num] = fallback
        return fallback
    try:
        soup = BeautifulSoup(r.text, "lxml")
        found: Optional[str] = None
        for a in soup.find_all("a", href=True):
            href = (a.get("href") or "").strip()
            if not href or "hymnary/audio/GG2013" not in href or ".mp3" not in href:
                continue
            raw = href.split("?")[0]
            if not raw.endswith(".mp3"):
                continue
            if raw.startswith("/"):
                found = "https://hymnary.org" + raw
            else:
                parsed = urlparse(raw)
                found = urlunparse((parsed.scheme, parsed.netloc, parsed.path, "", "", ""))
            break
        url = found or fallback
    except Exception as e:
        logger.debug("resolve_hymnary_audio_url parse %s: %s", page_url, e)
        url = fallback
    _hymnary_audio_resolve_cache[num] = url
    if url != fallback:
        logger.debug("resolve_hymnary_audio_url number=%s -> %s", num, url)
    return url


def generate_liturgy(
    *,
    occasion: str,
    scriptures: List[str],
    hymns: List[Dict[str, str]],
    sections: List[str],
    api_key: Optional[str] = None,
    user_overrides: Optional[Dict[str, str]] = None,
) -> Dict[str, str]:
    """
    Use OpenAI to generate liturgy text for the requested sections.
    If user_overrides[section] is non-empty, that text is used instead of generating.
    Returns dict mapping section key -> plain text.
    """
    overrides = user_overrides or {}
    client = None
    if OpenAI:
        key = api_key or os.getenv("OPENAI_API_KEY")
        if key:
            client = OpenAI(api_key=key)

    if not client:
        return {
            s: f"[Configure OPENAI_API_KEY to generate {s.replace('_', ' ')}.]"
            for s in sections
        }

    hymn_lines = "\n".join(
        f"- {h.get('title', '')} (GG2013 #{h.get('number', '')})" for h in hymns
    )
    scripture_lines = "\n".join(f"- {s}" for s in scriptures) if scriptures else "None specified."

    system = (
        "You are a thoughtful worship writer for Christian liturgy from a moderate Reformed perspective, "
        "in line with PC(USA) theology. Write in clear, inclusive language. "
        "Avoid exclusively male references to God: use 'God' by name, or varied language (e.g. 'God who is Father, Son, and Holy Spirit' when trinitarian language fits); "
        "do not use only 'he/him/his' or 'Lord' alone for God; you may use 'Lord' as one among other titles. "
        "Keep each piece concise and usable in worship. Output only the liturgy text, no meta-commentary or labels."
    )

    out = {}
    for section in sections:
        if overrides.get(section, "").strip():
            out[section] = overrides[section].strip()
            continue
        if section == "call_to_worship":
            prompt = (
                f"Write a Call to Worship for: {occasion}. "
                f"Scriptures: {scripture_lines}. Opening hymn: {hymns[0].get('title', '') if hymns else 'N/A'}. "
                "Use exactly this format with four parts: 'Leader: ' (2-3 lines), then 'People: ' (one short response), "
                "then 'Leader: ' again (2-3 lines), then 'People: ' again (one short response)."
            )
        elif section == "prayer_of_confession":
            prompt = (
                f"Write a Prayer of Confession for: {occasion}. "
                f"Scriptures: {scripture_lines}. "
                "One short paragraph. First person plural (we). End with a line inviting silence or a brief moment of confession."
            )
        elif section == "assurance":
            prompt = (
                f"Write only the Leader line for Assurance of Pardon for: {occasion}, "
                "grounded in God’s grace in Christ. One sentence. Do not include "
                "'People:' or 'Thanks be to God'—that will be added separately. "
                "Start your response with 'Leader: ' followed by the sentence."
            )
        elif section == "opening_prayer":
            prompt = (
                f"Write an Opening Prayer (collect) for: {occasion}. "
                f"Scriptures: {scripture_lines}. "
                "Around 100 words. Address God, thank or praise, and ask for one thing fitting the day. End with 'Amen.'"
            )
        elif section == "prayer_for_illumination":
            prompt = (
                f"Write a Prayer for Illumination for: {occasion}. "
                f"Scriptures: {scripture_lines}. "
                "Write 3-5 sentences asking God to open hearts and minds to the Scripture, "
                "that we may hear and respond. End with 'Amen.'"
            )
        elif section == "offertory_prayer":
            prompt = (
                f"Write a brief Offertory Prayer for: {occasion}. "
                "One or two sentences dedicating our gifts and ourselves to God's service. End with 'Amen.'"
            )
        elif section == "prayers_of_the_people":
            prompt = (
                f"Write Prayers of the People for: {occasion}. "
                f"Scriptures: {scripture_lines}. Hymns: {hymn_lines}. "
                "Write a substantial, full prayer (at least 10–15 paragraphs) in 'out to in' order: "
                "First, prayers for the world (nations, creation, peace, the suffering). "
                "Second, prayers for the Church universal and our denomination and congregation. "
                "Third, prayers for our country and our local community and leaders. "
                "Fourth, prayers for ourselves and our families. "
                "Then include an explicit invitation for the congregation to share joys and concerns aloud "
                "(e.g. 'Let us now lift up the joys and concerns of this congregation' or 'You are invited to name aloud...'). "
                "Then include a clear bid for a moment of silence—to lift up individuals or situations, or simply to sit in silence before God. "
                "Use 'we pray,' 'let us pray,' or similar. End with a closing that leads into the Lord's Prayer or a final amen. "
                "Write in full sentences and paragraphs; this should feel like a complete, unhurried pastoral prayer."
            )
        elif section == "benediction":
            prompt = (
                f"Write a Benediction (1–3 sentences) for: {occasion}. "
                f"Scriptures: {scripture_lines}. "
                "Send the people out to serve and share God’s love. You may invoke the Trinity. End with 'Amen.'"
            )
        else:
            out[section] = ""
            continue

        model = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
        try:
            r = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=1024,
            )
            text = (r.choices[0].message.content or "").strip()
            out[section] = text
        except Exception as e:
            out[section] = f"[Error generating {section}: {e}]"

    return out


def _add_custom_elements_after(
    doc,
    anchor: str,
    custom_elements: List[Dict[str, Any]],
) -> None:
    """Add any custom elements that are inserted after this anchor."""
    for ce in custom_elements:
        if ce.get("insert_after") == anchor and ce.get("label"):
            doc.add_paragraph(ce["label"], style="Heading 2")
            if ce.get("text"):
                doc.add_paragraph(ce["text"])
            doc.add_paragraph()


def build_docx(
    *,
    occasion: str,
    date: str,
    scriptures: List[str],
    hymns: List[Dict[str, str]],
    liturgy: Dict[str, str],
    include_placeholders: bool = True,
    sermon_title: Optional[str] = None,
    selected_ot_ref: Optional[str] = None,
    selected_nt_ref: Optional[str] = None,
    scripture_full_texts: Optional[Dict[str, str]] = None,
    include_sermon: bool = True,
    include_prayers_of_the_people: bool = True,
    include_communion: bool = False,
    custom_elements: Optional[List[Dict[str, Any]]] = None,
) -> BytesIO:
    """
    Build a Word document with the worship service order and generated liturgy.
    include_sermon: include Sermon Title section (for pastor copy; omit for secretary).
    include_prayers_of_the_people: include Prayers of the People (for pastor copy; omit for secretary).
    include_communion: include The Sacrament of the Lord's Supper liturgy (e.g. first Sunday of month).
    Returns a BytesIO buffer containing the .docx.
    """
    if not Document:
        raise RuntimeError("python-docx is required. pip install python-docx")

    custom = custom_elements or []
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Worship Service\n{occasion}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(date).font.size = Pt(12)
    doc.add_paragraph()

    # 1. Call to Worship (Leader / People, People in bold)
    if liturgy.get("call_to_worship"):
        doc.add_paragraph("Call to Worship", style="Heading 2")
        _add_leader_people_paragraph(doc, liturgy["call_to_worship"])
        doc.add_paragraph()
    _add_custom_elements_after(doc, "call_to_worship", custom)

    # 2. Opening Prayer
    if liturgy.get("opening_prayer"):
        doc.add_paragraph("Opening Prayer", style="Heading 2")
        doc.add_paragraph(liturgy["opening_prayer"])
        doc.add_paragraph()
    _add_custom_elements_after(doc, "opening_prayer", custom)

    # 3. First Hymn
    if hymns:
        doc.add_paragraph("First Hymn", style="Heading 2")
        h = hymns[0]
        doc.add_paragraph(f"{h.get('title', '')} — #{h.get('number', '')}")
        doc.add_paragraph()
    _add_custom_elements_after(doc, "first_hymn", custom)

    # 4. Prayer of Confession (bold)
    if liturgy.get("prayer_of_confession"):
        doc.add_paragraph("Prayer of Confession", style="Heading 2")
        p = doc.add_paragraph()
        p.add_run(liturgy["prayer_of_confession"]).bold = True
        doc.add_paragraph()
    _add_custom_elements_after(doc, "prayer_of_confession", custom)

    # 5. Assurance of Pardon (Leader: ... / People: Thanks be to God! Amen. in bold)
    if liturgy.get("assurance"):
        doc.add_paragraph("Assurance of Pardon", style="Heading 2")
        _add_assurance_paragraph(doc, liturgy["assurance"])
        doc.add_paragraph()
    _add_custom_elements_after(doc, "assurance", custom)

    # 6. Prayer for Illumination
    if liturgy.get("prayer_for_illumination"):
        doc.add_paragraph("Prayer for Illumination", style="Heading 2")
        doc.add_paragraph(liturgy["prayer_for_illumination"])
        doc.add_paragraph()
    _add_custom_elements_after(doc, "prayer_for_illumination", custom)

    # 7. Old Testament Reading (reference only)
    ot_ref = selected_ot_ref or (scriptures[0] if scriptures else None)
    if ot_ref:
        doc.add_paragraph("Old Testament Reading", style="Heading 2")
        doc.add_paragraph(ot_ref)
        doc.add_paragraph()
    _add_custom_elements_after(doc, "ot_reading", custom)

    # 8. New Testament Reading (reference only)
    nt_ref = selected_nt_ref or (scriptures[1] if len(scriptures) > 1 else None)
    if nt_ref:
        doc.add_paragraph("New Testament Reading", style="Heading 2")
        doc.add_paragraph(nt_ref)
        doc.add_paragraph()
    _add_custom_elements_after(doc, "nt_reading", custom)

    # 9. Sermon Title (optional; for pastor copy only)
    if include_sermon:
        doc.add_paragraph("Sermon Title", style="Heading 2")
        doc.add_paragraph(sermon_title.strip() if sermon_title and sermon_title.strip() else "[Sermon title]")
        doc.add_paragraph()
    _add_custom_elements_after(doc, "sermon", custom)

    # 10. Affirmation of Faith
    doc.add_paragraph("Affirmation of Faith", style="Heading 2")
    doc.add_paragraph("Apostles' Creed (or as printed)")
    doc.add_paragraph()
    _add_custom_elements_after(doc, "affirmation_of_faith", custom)

    # 11. Second Hymn
    if len(hymns) > 1:
        doc.add_paragraph("Second Hymn", style="Heading 2")
        h = hymns[1]
        doc.add_paragraph(f"{h.get('title', '')} — #{h.get('number', '')}")
        doc.add_paragraph()
    _add_custom_elements_after(doc, "second_hymn", custom)

    # 11b. Communion liturgy (after second hymn when communion is included)
    if include_communion:
        _add_communion_liturgy(doc)
    _add_custom_elements_after(doc, "communion", custom)

    # 12. Prayers of the People (optional; for pastor copy only)
    if include_prayers_of_the_people and liturgy.get("prayers_of_the_people"):
        doc.add_paragraph("Prayers of the People", style="Heading 2")
        doc.add_paragraph(liturgy["prayers_of_the_people"])
        doc.add_paragraph()
    _add_custom_elements_after(doc, "prayers_of_the_people", custom)

    # 13. Offertory Prayer
    if liturgy.get("offertory_prayer"):
        doc.add_paragraph("Offertory Prayer", style="Heading 2")
        doc.add_paragraph(liturgy["offertory_prayer"])
        doc.add_paragraph()
    _add_custom_elements_after(doc, "offertory_prayer", custom)

    # 14. Third Hymn
    if len(hymns) > 2:
        doc.add_paragraph("Third Hymn", style="Heading 2")
        h = hymns[2]
        doc.add_paragraph(f"{h.get('title', '')} — #{h.get('number', '')}")
        doc.add_paragraph()
    _add_custom_elements_after(doc, "third_hymn", custom)
    _add_custom_elements_after(doc, "benediction", custom)  # "Before Benediction"

    # Benediction
    if liturgy.get("benediction"):
        doc.add_paragraph("Benediction", style="Heading 2")
        doc.add_paragraph(liturgy["benediction"])

    _add_custom_elements_after(doc, "end", custom)  # At the end (after Benediction)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
